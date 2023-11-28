import re
from docx import Document


class appFuncs():
    def __init__(self) -> None:
        self.doc = Document('models/Contrato de Consignação.docx')
        self.var_init = '<'
        self.var_end = '>'
        self.vars = self._find_variables('<', '>')
        self.var_dict = self._set_data_dic()
        self._format_doc()

    def _find_variables(self, var_sign_start, var_sign_end):
        pattern = re.compile(
            fr'{re.escape(var_sign_start)}(.*?){re.escape(var_sign_end)}')
        found_varibles = []
        for paragraph in self.doc.paragraphs:
            corresp = pattern.findall(paragraph.text)
            found_varibles.extend(corresp)
        return list(set(found_varibles))

    def _set_data_dic(self):
        var_dict = {}
        for var in self.vars:
            data = input(f'{var} : ')
            var_dict[var] = data
        return var_dict

    def _format_doc(self):
        for paragraph in self.doc.paragraphs:
            for var, data in self.var_dict.items():
                marcador = f"{self.var_init}{var}{self.var_end}"
                paragraph.text = paragraph.text.replace(marcador, data)
        # Salvar o documento modificado
        self.doc.save('out/Contrato_Modificado.docx')


main = appFuncs()

print(main.var_dict)
